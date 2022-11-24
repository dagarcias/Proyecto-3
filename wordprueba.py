from docx import Document

# Creación del documento
document = Document()

document.add_paragraph('binario de la letra')
document.add_paragraph('binario de la letra')

'''
While no mas caracteres:
    if letra ya esta usada:
        document.add_paragraph('ubi arbol de letra ')
    else:
        document.add_paragraph('ubi padre de la letra')
        document.add_paragraph('binario de la nueva letra')

'''
document.save('ejemplo.docx')