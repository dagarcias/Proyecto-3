import ArbolDeHuffman
from Pruebas2 import switchL
from docx import Document
from binario import diccionario
document = Document()

class node:

    def __init__(self, frecuencia, enumeracion):
        self.izq = None
        self.der = None
        self.posicion = None
        self.frecuencia = frecuencia
        self.simbolos = None
        self.enumeracion = enumeracion
        
    def get_posicion(self):
        return self.posicion
    
    def set_posicion(self, posicion):
        self.posicion = posicion

    def get_enumeracion(self):
        return self.enumeracion

    def get_frecuencia(self):
        return self.frecuencia

    def set_izq(self, hijoizq):
        self.izq = hijoizq

    def set_der(self, hijoder):
        self.der = hijoder

    def get_izq(self):
        return self.izq

    def get_der(self):
        return self.der

    def set_simbolo(self, simbolo):
        self.simbolos = simbolo

    def get_simbolo(self):
        return self.simbolos

class arbol:
    def __init__(self, raiz):
        self.raiz = raiz

lista = ArbolDeHuffman.main()
listaFrecuencia = ArbolDeHuffman.returnFrecuencias()
lista.reverse()
n = 107

resta = 1
nodoUsado = [1]

def crearArbol():
    global resta
    raiz = node(lista[0][2], 2*n-resta)
    resta += 1   
    raiz.set_izq(node(lista[0][0], 2*n-resta))
    resta += 1
    raiz.set_der(node(lista[0][1], 2*n-resta))
    
    for x in range(0,len(lista)-1):
        nodoUsado.append(0)  
    
    crearNodos(raiz.get_izq(), lista[0][0], "1")  
    crearNodos(raiz.get_der(), lista[0][1], "0")
    return raiz

def crearNodos(nodo, valor, posicion):
    global resta
    indiceItem = 0
    nodo.set_posicion(posicion)
    for item in lista:
        if item[2] == valor and nodoUsado[indiceItem] == 0:
            nodoUsado[indiceItem] = 1    
            
            resta += 1
            nodo.set_izq(node(item[0], 2*n-resta))
            crearNodos(nodo.get_izq(), item[0], posicion+"1")
            #posicion = posicion[:-1]
            
            resta += 1
            nodo.set_der(node(item[1], 2*n-resta))
            crearNodos(nodo.get_der(), item[1], posicion+"0")
            break
        indiceItem += 1
         

nodosHojas = []
def hallarHojas(nodo): 
    if nodo.get_izq() == None and nodo.get_der() == None:
        nodosHojas.append(nodo)
    else:
        hallarHojas(nodo.get_izq())
        hallarHojas(nodo.get_der())
    return nodosHojas
        
def infijo(nodo):
    if nodo != None:
        infijo(nodo.get_izq())
        print(
            f"[{nodo.get_simbolo(), nodo.get_frecuencia(), nodo.get_enumeracion(), nodo.get_posicion()}]")
        infijo(nodo.get_der())

raiz = crearArbol()
listaHojas = hallarHojas(raiz)

def añadirSimbolos(listaHojas, listaOrden):
    hojasUsadas = []
    for x in listaOrden:
        hojasUsadas.append(0)
    
    for hoja in listaHojas:
        final = False
        indice = 0
        while not final:
            if listaOrden[indice][1] == hoja.get_frecuencia():
                hoja.set_simbolo(listaOrden[indice][0])
                listaOrden.pop(indice)
                final = True
            indice += 1

listaOrden = switchL(listaFrecuencia)
añadirSimbolos(listaHojas, listaOrden)
infijo(raiz)

datosHojas = []
for i in listaHojas:
    datosHojas.append(i.get_frecuencia())
   
posEncontrada = None
def posicionLetra(letra, nodo):
    global posEncontrada
    if nodo!=None:
        if nodo.get_simbolo() == letra:
            posEncontrada = nodo.get_posicion()
        else:
            posicionLetra(letra, nodo.get_izq())
            posicionLetra(letra, nodo.get_der())                
    
txt = ArbolDeHuffman.get_txt()

listaLetras = []
for letra in txt:
    posicionLetra(letra, raiz)
    listaLetras.append(posEncontrada)

print(listaLetras)

print(f"HOJAS: {datosHojas}")

posPadre = None
def hallarPosPadre(letra, nodo):
    global posPadre
    
    if nodo != None:
        if nodo.get_izq() != None and nodo.get_der() != None: 
            nodoDer = nodo.get_der().get_simbolo()
            nodoIzq = nodo.get_izq().get_simbolo()
              
            if nodoDer == letra or nodoIzq == letra:
                posPadre = nodo.get_posicion() 
            else:
                hallarPosPadre(letra, nodo.get_izq())
                hallarPosPadre(letra, nodo.get_der())
         
        
        
primera=True
recorridoCompleto = False
letrasUsadas = []

for letra in txt:
    if primera:
        document.add_paragraph(f"{letra}: {diccionario[letra]}")
        primera=False
    else:   
        if letra in letrasUsadas:
            posicionLetra(letra, raiz)   
            document.add_paragraph(f"{letra}: {posEncontrada}")
        else:
            hallarPosPadre(letra, raiz)
            letrasUsadas.append(letra)  
            document.add_paragraph(posPadre)
            document.add_paragraph(f"{letra}: {diccionario[letra]}")

document.save('recorrido.docx')

input('enter para salir: ')


